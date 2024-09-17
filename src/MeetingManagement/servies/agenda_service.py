import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader
import docx
import json
from MeetingManagement import logger
from langchain_core.documents import Document
from langchain.vectorstores import Chroma
from langchain_google_genai import GoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.prompts import PromptTemplate
import re
from uuid import uuid4

load_dotenv()

def get_cleaned_text(text):

    text = re.sub(r'[\n\t\v\r\f]', ' ', text)

    text = re.sub(r'[^\x20-\x7E]', '', text)

    text = re.sub(r'\s+', ' ', text)

    text = text.strip()

    return text


def extract_text_from_documents():

    try:
        if os.path.exists("database/documents/original"):

            document_path_list = os.listdir("database/documents/original")
            print(document_path_list)

            logger.info("Documents found in the database.")

    except FileNotFoundError as e:
        logger.error(f"{e}\n")

    text = ""

    for file in document_path_list:
        file_path = os.path.join("database/documents/original", file)

        if file.endswith('.pdf'):
            pdf = PdfReader(file_path)
            for page in pdf.pages:
                text += page.extract_text() + '\n'
            text = get_cleaned_text(text)
        
        elif file.endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'
            text = get_cleaned_text(text)

        elif file.endswith('.txt'):
            with open(file_path, 'r', encoding = 'utf-8') as f:
                text += f.read()
            text = get_cleaned_text(text)

        else:
            raise ValueError("Unsupported file format")
        
    if not os.path.exists("database/documents/processed"):
        os.makedirs("database/documents/processed")

    with open("database/documents/processed/preprocessed.txt", 'w') as f:
        f.write(text) 

    return text


def split_document_content(words_per_chunk = 50):

    text = extract_text_from_documents()

    words = text.split()

    text_chunks = []

    for i in range(0, len(text), words_per_chunk):
        text_chunks.append(' '.join(words[i:i+words_per_chunk]))

    documents = [Document(page_content = content) for content in text_chunks]

    return documents


def intialize_vector_stores(documents):
    
    embeddings = HuggingFaceEmbeddings()

    db = Chroma(embedding_function = embeddings)

    uuids = [str(uuid4()) for _ in range(len(documents))]
    db.add_documents(documents)

    return db


def get_agenda_text(discussion_point, vector_store):
    
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    genai = GoogleGenerativeAI(model = "gemini-pro", api_key = gemini_api_key)

    question_text = ""

    for point in discussion_point:

        similar_docs = vector_store.similarity_search(point, k = 2)

        question_text += f"Discussion Point: {point}\n"
        question_text += "Relevant Document Excerpts:\n"

        for doc in similar_docs:
            question_text += f" - {doc.page_content}\n"
        
        question_text += "\n"

    prompt_text = PromptTemplate.from_template(
        """
            You are a member of the meeting management team. You have been tasked with generating agenda items for the upcoming meeting.
            Provide a very clear and well defined agenda item based on the following discussion points and relevant document excerpts:

            Discussion points and their relevant document excerpts: {question_text}

            Remeber it is important to provide a clear title for the agenda item and key points to be discussed.
            Also it should be in natural language and easy to understand.
            No need to provide any additional information just steek with the provided information.
        """
    )

    chain = prompt_text | genai

    agenda_text = chain.invoke({"question_text" : question_text})

    return agenda_text


def agenda_generation():
    """
        main function to generate agenda items from the document
    """

    documents = split_document_content()

    discussion_points_file_path = "database/discussion_points/discussion_points.json"
    index_name = "chromaDB-meeting-agenda"
    
    with open(discussion_points_file_path, 'r') as f:
        discussion_points = json.load(f)

    vector_store = intialize_vector_stores(documents = documents)

    agenda_text = get_agenda_text(discussion_points, vector_store)

    return agenda_text