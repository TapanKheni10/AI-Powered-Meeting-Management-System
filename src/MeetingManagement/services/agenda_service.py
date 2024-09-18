import os
import docx
import json
import string
from uuid import uuid4
from PyPDF2 import PdfReader
from dotenv import load_dotenv

from langchain.vectorstores import Chroma
from langchain_core.documents import Document
from langchain.prompts import PromptTemplate
from langchain_google_genai import GoogleGenerativeAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from MeetingManagement import logger
from MeetingManagement.constants import DOCUMENTS_DIR, DISCUSSION_POINTS_DIR

load_dotenv()

def get_cleaned_text(text) -> str:
    """
    Clean the input text by removing non-printable characters and extra spaces.

    Args:
        text (str): The input text to be cleaned.

    Returns:
        str: The cleaned text.
    """

    # Remove non-ASCII characters and replace whitespaces with a single space
    text = ''.join([ch if ch in string.printable else ' ' for ch in text])

    # Remove extra spaces and string leading / trailing spaces
    return ' '.join(text.split())


def extract_text_from_documents() -> str:
    """
    Extract text from various document types for the user uploaded documents.

    Returns:
        str: Concatenated text from all documents.
    """

    try:
        original_dir_path = os.path.join(DOCUMENTS_DIR, "original")
        if os.path.exists(original_dir_path):
            document_path_list = os.listdir(original_dir_path)
            logger.info("Documents found in the database.")
        else:
            logger.error("Document directory not found.")

    except FileNotFoundError as e:
        logger.error(f"Error accessing document directory: {e}")

    text = ""

    logger.info("Extracting text from the uploaded documents...")

    # Extract text from each document
    for file in document_path_list:
        file_path = os.path.join(original_dir_path, file)

        # Extract text from PDF files
        if file.endswith('.pdf'):
            pdf = PdfReader(file_path)
            for page in pdf.pages:
                text += page.extract_text() + '\n'
        
        # Extract text from DOCX files
        elif file.endswith('.docx'):
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + '\n'

        # Extract text from TXT files
        elif file.endswith('.txt'):
            with open(file_path, 'r', encoding = 'utf-8') as f:
                text += f.read()

        else:
            raise ValueError("Unsupported file format")
        
    logger.info("Text extraction complete.")
        
    text = get_cleaned_text(text)
    logger.info("Text cleaning complete.")
        
    # Save the preprocessed text to a file
    processed_dir_path = os.path.join(DOCUMENTS_DIR, "processed")
    os.makedirs(processed_dir_path, exist_ok = True)

    with open(os.path.join(processed_dir_path, "processed.txt"), 'w') as f:
        f.write(text) 

    return text


def get_text_chunks(words_per_chunk = 100):
    """
    Split the extracted text into chunks.

    Args:
        words_per_chunk (int): Number of words per chunk. Default is 100.

    Returns:
        list: List of text chunks.
    """

    # Extract text from the uploaded documents
    text = extract_text_from_documents()

    # Split the text into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size = words_per_chunk, chunk_overlap = 30)
    text_chunks = text_splitter.split_text(text)

    logger.info(f"Text split into {len(text_chunks)} chunks.")

    return text_chunks


def get_vector_store(documents):
    """
    Create a vector store from the given documents.

    Args:
        documents (list): List of Document objects.

    Returns:
        Chroma (object): A Chroma vector store containing the document embeddings.
    """

    logger.info("Creating vector store...")

    # Initialize the Hugging Face embeddings
    embeddings = HuggingFaceEmbeddings()

    # Initialize the Chroma vector store
    chroma_db = Chroma(embedding_function = embeddings)

    logger.info("Adding documents to the vector store...")
    # Add the documents to the vector store
    uuids = [str(uuid4()) for _ in range(len(documents))]
    chroma_db.add_documents(documents, ids = uuids)

    return chroma_db

def get_conversational_chain():
    """
    Create a conversational chain using the Gemini API.

    Returns:
        chain (object): A conversational chain object.
    """

    # Initialize the Google Generative AI
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    genai = GoogleGenerativeAI(model = "gemini-pro", api_key = gemini_api_key)

    # Create a conversational chain
    prompt = PromptTemplate.from_template(
        """
            You are a member of the meeting management team. You have been tasked with generating agenda items for the upcoming meeting.
            Provide a very clear and well defined agenda item based on the following discussion points and relevant document excerpts:

            Discussion points and their relevant document excerpts: {question_text}

            Remeber it is important to provide a clear title for the agenda item and key points to be discussed.
            Also it should be in natural language and easy to understand.
            No need to provide any additional information just steek with the provided information.
        """
    )

    chain = prompt | genai

    logger.info("Conversational chain created.")

    return chain


def get_agenda(discussion_point, vector_store):
    """
    Generate agenda items based on discussion points and relevant document excerpts.

    Args:
        discussion_points (list): List of discussion points.
        vector_store (Chroma): The vector store containing document embeddings.

    Returns:
        str: Generated agenda items.
    """
    
    agenda_list = []

    logger.info("Finding relevant document excerpts for each discussion point...")
    # Get the most relevant document excerpts for each discussion point
    for point in discussion_point:

        related_documents = vector_store.similarity_search(point, k = 2)

        agenda_list.append(f"Point: {point}\nRelevant Document Excerpts:\n")

        for i, doc in enumerate(related_documents):
            agenda_list.append(f"{i}) {doc.page_content}")
        
        agenda_list.append("\n")

    agenda_summary_text = '\n'.join(agenda_list)

    # Generate the agenda items using the conversational chain
    chain = get_conversational_chain()

    logger.info("Generating agenda items using the conversational chain...")

    response = chain.invoke({"question_text": agenda_summary_text})

    return response


def agenda_generation():
    """
    Main function to generate agenda items from the documents.

    Returns:
        str: Generated agenda items.
    """

    logger.info("Generating agenda items...")

    text_chunks = get_text_chunks()

    # Create a list of Document objects
    documents = [Document(page_content = chunk) for chunk in text_chunks]
    logger.info(f"Number of documents created: {len(documents)}")

    discussion_points_file_path = os.path.join(DISCUSSION_POINTS_DIR, 'discussion_points.json')
    
    # Load the discussion points from the JSON file
    with open(discussion_points_file_path, 'r') as f:
        discussion_points = json.load(f)

    vector_store = get_vector_store(documents = documents)

    logger.info("Vector store created.")
    logger.info("Generating agenda items...")

    response = get_agenda(discussion_points, vector_store)

    logger.info("Agenda items generated successfully.")

    return response